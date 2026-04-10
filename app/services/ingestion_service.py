# app/services/ingestion_service.py

import logging
from io import BytesIO

from llama_index.core import Document, VectorStoreIndex
from llama_index.core.node_parser import SimpleNodeParser

from app.core.vector_store import storage_context
from app.utils.text_cleaner import clean_text  # 🔥 IMPORTANT

from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)


# ================= TEXT EXTRACTION =================
def extract_text(file, filename: str):
    logger.info(f"📄 Extracting text from file: {filename}")

    file.seek(0)
    content = file.read()

    logger.info(f"📦 File size: {len(content)} bytes")

    try:
        if filename.endswith(".pdf"):
            logger.info("📘 Processing PDF file")
            reader = PdfReader(BytesIO(content))
            text = ""

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                logger.info(f"📄 Page {i} length: {len(page_text)}")
                text += page_text + "\n"

            return text

        elif filename.endswith(".docx"):
            logger.info("📘 Processing DOCX file")
            doc = DocxDocument(BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs])

        elif filename.endswith(".xlsx"):
            logger.info("📊 Processing Excel file")
            df = pd.read_excel(BytesIO(content))
            return df.to_string()

        else:
            logger.info("📄 Processing text-based file")
            return content.decode("utf-8", errors="ignore")

    except Exception as e:
        logger.error("❌ Error during text extraction")
        logger.exception(e)
        return ""


# ================= MAIN INGESTION =================
async def process_file(file, user_id):
    try:
        logger.info("\n" + "=" * 60)
        logger.info(f"🚀 START INGESTION | user_id: {user_id}")
        logger.info("=" * 60)

        # ---------------- EXTRACT TEXT ----------------
        raw_text = extract_text(file.file, file.filename)

        logger.info(f"📏 Raw text length: {len(raw_text)}")

        if not raw_text.strip():
            logger.warning("⚠️ Empty document detected after extraction")
            return 0

        # ---------------- CLEAN TEXT (🔥 VERY IMPORTANT) ----------------
        text = clean_text(raw_text)

        logger.info(f"🧹 Cleaned text length: {len(text)}")

        if not text.strip():
            logger.warning("⚠️ Empty document after cleaning")
            return 0

        # ---------------- DOCUMENT ----------------
        logger.info("🧱 Creating Document object")

        metadata = {
            "user_id": str(user_id),
            "filename": file.filename
        }

        document = Document(
            text=text,
            metadata=metadata
        )

        logger.info(f"📄 Document created with metadata: {metadata}")

        # ---------------- CHUNKING ----------------
        logger.info("🔪 Splitting into nodes")

        parser = SimpleNodeParser.from_defaults(
            chunk_size=400,        # 🔥 slightly improved
            chunk_overlap=40
        )

        nodes = parser.get_nodes_from_documents([document])

        logger.info(f"📦 Total nodes created: {len(nodes)}")

        # 🔥 DEBUG NODES (LIMITED)
        for i, node in enumerate(nodes[:3]):  # show only first 3
            logger.info(f"--- NODE {i} ---")
            logger.info(f"Text preview: {node.text[:150]}")
            logger.info(f"Metadata: {node.metadata}")

        # ---------------- VECTOR STORE ----------------
        logger.info("💾 Storing nodes into Vector DB")

        index = VectorStoreIndex(
            nodes,
            storage_context=storage_context
        )

        index.storage_context.persist()

        logger.info("✅ Data stored successfully in vector DB")

        logger.info("=" * 60)
        logger.info("🏁 INGESTION COMPLETED")
        logger.info("=" * 60)

        return len(nodes)

    except Exception as e:
        logger.error("❌ INGESTION FAILED")
        logger.exception(e)  # 🔥 full traceback
        return f"Error: {str(e)}"